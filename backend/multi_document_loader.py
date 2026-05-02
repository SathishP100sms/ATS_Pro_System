import os
import fitz
from docx import Document



def load_document(file_path: str) -> str:
    """
    Load text from PDF, DOCX, or TXT file.
    Returns cleaned text.
    """

    if not os.path.exists(file_path):
        raise FileNotFoundError("File not found")

    if os.path.getsize(file_path) == 0:
        raise ValueError("File is empty")

    file_extension = os.path.splitext(file_path)[1].lower()

    if file_extension == ".pdf":
        return _load_pdf(file_path)

    elif file_extension == ".docx":
        return _load_docx(file_path)

    elif file_extension == ".txt":
        return _load_txt(file_path)

    else:
        raise ValueError("Unsupported file type")


# ---------------- PDF ---------------- #
def _load_pdf(file_path: str) -> str:
    try:
        text_parts = []

        with fitz.open(file_path) as doc:
            for page in doc:
                text_parts.append(page.get_text())

        return _clean_text("\n".join(text_parts))

    except Exception as e:
        raise Exception(f"Error processing PDF file: {str(e)}")


# ---------------- DOCX ---------------- #
def _load_docx(file_path: str) -> str:
    try:
        doc = Document(file_path)
        text_parts = []

        # Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Tables 
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))

        return _clean_text("\n".join(text_parts))

    except Exception as e:
        raise Exception(f"Error processing DOCX file: {str(e)}")


# ---------------- TXT ---------------- #
def _load_txt(file_path: str) -> str:
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return _clean_text(f.read())

    except Exception as e:
        raise Exception(f"Error processing TXT file: {str(e)}")


# ---------------- CLEANING ---------------- #
def _clean_text(text: str) -> str:
    """
    Basic cleaning to improve embedding + keyword extraction
    """
    if not text:
        return ""

    # Remove excessive whitespace
    text = " ".join(text.split())

    # Normalize newlines
    text = text.replace("\n", " ").replace("\r", " ")

    return text.strip()
    
  

